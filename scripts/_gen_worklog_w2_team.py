"""
Generate Individual Work Logs for 3 team members — WEEK 2
Based on actual project progress (19/05/2026 - 23/05/2026)
Source: covertV1.md (tail), covertV2.md, convert2F.md, current conversation
"""
import docx

TEMPLATE = r'd:\OJT\Worklog\Worklog - Wx - Student ID - FullName.docx'
PROJECT_NAME = "ThuanPhongNhi: Uncertainty-Aware Federated Learning for Emotion Recognition in Conversations"

members = [
    {
        "name": "Đinh Đại Lộc",
        "id": "22110046",
        "filename": r"d:\OJT\Worklog\Worklog - W2 - 22110046 - DinhDaiLoc.docx",
        "tasks": [
            ("Chạy P1 (EAFA Extreme): 80 experiments noise robustness (noise 30-50%, alpha 0.1-0.3) trên MELD + IEMOCAP",
             "Completed", "4h",
             "EAFA vượt FedAvg +2.37% ở noise=50% (p<0.01). results_eafa_extreme.json"),
            ("Fix P2 DailyDialog baselines: sửa CoMPM và SPCL từ dummy model thành real architecture (Transformer encoder, Prototypical CL)",
             "Completed", "5h",
             "CoMPM dùng TransformerEncoder, SPCL dùng PrototypicalCL. 75/75 experiments xong"),
            ("ECR Hyperparameter Tuning: 2-stage (48 coarse + 45 validate) tìm optimal sigma, lambda, ramp_start",
             "Completed", "3h",
             "Best config: sigma=0.01, lambda=0.3, ramp=3. Delta vs default: +0.05%"),
            ("ECR Augmentation Test: so sánh 7 augmentation strategies (Gaussian noise vs Feature Dropout)",
             "In Progress", "2h",
             "scripts/run_ecr_augtest.py — 63 experiments, feature dropout có thể tạo perturbation mạnh hơn"),
            ("Comprehensive Survey: phân tích chi tiết 17 papers, so sánh SOTA, xác định 4 research gaps",
             "Completed", "3h",
             "docs/comprehensive_survey.md — 4 gaps validated, AAAI roadmap tạo xong"),
            ("Xác minh literature survey: kiểm tra recency (≥2023), đề xuất bổ sung 5 papers mới (InstructERC, SPCL, CoMPM, Semi-FedSER, LineConGAT)",
             "Completed", "2h",
             "survey_verification.md — chỉ 3/13 papers ≥2023, cần bổ sung 5 papers"),
        ],
        "total_time": "~19h",
        "summary_tasks": """-Key tasks done:
1. Hoàn thành P1 Extreme Experiments (80 runs): EAFA vượt FedAvg +2.37% tại noise=50% với p<0.01
2. Fix P2 DailyDialog: sửa CoMPM/SPCL từ dummy model thành real architectures, chạy xong 75/75 experiments
3. ECR Hyperparameter Tuning 2-stage: tìm optimal config (sigma=0.01, lambda=0.3, ramp=3)
4. Tạo comprehensive survey 17 papers, xác nhận 4 research gaps vẫn valid
5. Xác minh recency survey: đề xuất bổ sung 5 papers gần đây (2022-2025)
6. Khởi tạo ECR augmentation test (Feature Dropout vs Gaussian Noise)""",
        "summary_learned": """-Key things learned:
1. 2-stage hyperparameter search (coarse→validate) tiết kiệm 80% thời gian so với brute force
2. ECR hoạt động tốt nhất với perturbation nhẹ (sigma=0.01) + SSL weight thấp (lambda=0.3)
3. AAAI A* yêu cầu theoretical analysis — đây là bottleneck chính cho submission
4. Narrative paper quan trọng hơn số liệu: "competitive without threshold" > "beat by 0.1%"
5. Feature dropout có thể tạo learning signal mạnh hơn Gaussian noise cho Dirichlet KL""",
        "summary_literature": """-Literature read:
1. Lei et al. (2024) "InstructERC" — AAAI 2024 SOTA, LLM framework cho ERC (71.39% IEMOCAP)
2. Song et al. (2022) "SPCL" — Supervised Prototypical Contrastive Learning (69.74% IEMOCAP)
3. Lee & Lee (2022) "CoMPM" — Context Modeling with Speaker's Pre-trained Memory (~68% IEMOCAP)
4. Padi et al. (2023/2025) "LineConGAT" — Speaker-independent GNN cho ERC (76.50% IEMOCAP)
5. Ono & Wakaki (2025) "TED" — Turn Emphasis with Dialogue Feature Attention""",
        "summary_issues": """-Issues/Challenges:
1. ECR không cải thiện đáng kể khi tune hyperparameters — delta chỉ 0.05% trên MELD
2. DailyDialog ceiling effect: tất cả methods 87-88% WF1 → SSL không giúp thêm
3. IEMOCAP WF1 (58%) vẫn thấp hơn nhiều so với SOTA (71%) do RoBERTa-Base + 6-class
4. CoMPM/SPCL ban đầu là dummy models (dùng chung weights với FedAvg) — đã fix
5. Brute force tuning 432 experiments quá tốn thời gian — đã chuyển sang 2-stage approach""",
        "next_plan": """Plan for the next week:
1. Phân tích kết quả ECR augmentation test (Feature Dropout vs Gaussian)
2. Bắt đầu viết paper AAAI LaTeX draft
3. Convergence analysis cho EAFA (theoretical contribution)""",
        "next_tasks": [
            ("Phân tích ECR augmentation test + quyết định final config", "Week 3 (May 26)"),
            ("Viết paper: Method section (ECR + EAFA formulation)", "Week 3 (May 28)"),
            ("Convergence analysis cho EAFA dưới non-IID", "Week 3 (May 30)"),
        ],
    },
    {
        "name": "Trần Phi Học",
        "id": "22110024",
        "filename": r"d:\OJT\Worklog\Worklog - W2 - 22110024 - TranPhiHoc.docx",
        "tasks": [
            ("Chạy EAFA 5-Seeds Noise Robustness: nâng từ 3→5 seeds (60 experiments) trên MELD + IEMOCAP",
             "Completed", "5h",
             "MELD 20% noise: p=0.044 significant! EAFA luôn có std thấp hơn FedAvg"),
            ("Thu thập và tổ chức 17 papers vào thư mục docs/papers/: download từ arXiv, ACL Anthology, IEEE",
             "Completed", "3h",
             "17 PDFs organized: 8 ERC, 2 FL, 2 EDL, 1 SSL, 2 FL+Emotion, 2 Recent"),
            ("Chạy Statistical Analysis đầy đủ: paired t-test, Wilcoxon, Cohen's d, 95% CI cho tất cả experiments",
             "Completed", "2h",
             "ECR vs FM: 3/6 significant (MELD). EAFA vs FedAvg: 1/6 significant"),
            ("Extract text từ 17 PDFs bằng PyPDF2 để hỗ trợ phân tích survey tự động",
             "Completed", "1h",
             "docs/papers_txt/ — 17 text files extracted, 30k chars/file cap"),
            ("Hỗ trợ chạy P2 DailyDialog: monitor 75 experiments, kiểm tra data integrity",
             "Completed", "3h",
             "Verified: CoMPM/SPCL 0/15 identical với FedAvg → fix thành công"),
            ("Tạo bảng so sánh SOTA chi tiết: MELD + IEMOCAP + DailyDialog vs 8 baseline methods",
             "Completed", "2h",
             "literature_survey.md — bảng SOTA với encoder size, year, setting"),
        ],
        "total_time": "~16h",
        "summary_tasks": """-Key tasks done:
1. Hoàn thành EAFA 5-seeds noise robustness (60 experiments): MELD 20% noise p=0.044 significant
2. Thu thập và tổ chức 17 research papers vào thư mục có cấu trúc
3. Chạy đầy đủ statistical analysis: t-test, Wilcoxon, Cohen's d cho toàn bộ experiments
4. Extract text từ PDFs cho phân tích survey tự động
5. Hỗ trợ verify P2 DailyDialog data integrity
6. Tạo bảng SOTA comparison chi tiết cho paper""",
        "summary_learned": """-Key things learned:
1. Statistical testing: paired t-test vs Wilcoxon — paired t-test phù hợp hơn cho seed-based comparison
2. Cohen's d effect size: >0.8 là "large", EAFA noise 40% đạt d=5.45 — rất mạnh
3. AAAI yêu cầu 5+ seeds cho statistical significance — 3 seeds không đủ
4. Cách tổ chức literature survey: phân loại theo category (ERC, FL, EDL, SSL)
5. PyPDF2 extract text từ academic papers — hữu ích cho automated analysis""",
        "summary_literature": """-Literature read:
1. Shou et al. (2026) "Comprehensive Survey on Multi-modal Conversational Emotion Recognition" — ACM TOIS
2. Feng et al. (2022) "Semi-FedSER: Federated Semi-Supervised SER" — INTERSPEECH, closest FL+SSL+Emotion
3. Qiu et al. (2025) "FedDISC" — NeurIPS 2025, FL+MERC competitor (diffusion-based modality recovery)
4. FedEmoNet (2026) "Federated Speech Emotion Recognition with DP" — PLOS ONE
5. FedMultiEmo (2025) "Real-time Multimodal Emotion Recognition via FL" — edge device focus""",
        "summary_issues": """-Issues/Challenges:
1. EAFA MELD 40% noise: delta +2.77% nhưng p=0.106 do FedAvg variance cao (1 seed kém 0.5351)
2. IEMOCAP noise experiments: EAFA không significant trên bất kỳ condition nào
3. Statistical power thấp với 3 seeds — nâng lên 5 seeds giúp nhưng vẫn chưa đủ
4. PDF text extraction quality không đồng đều — scan-based PDFs không đọc được
5. Cần thêm papers 2023+ để đáp ứng reviewer yêu cầu recency""",
        "next_plan": """Plan for the next week:
1. Hỗ trợ viết paper: Dataset Description + Experimental Setup sections
2. Tạo figures/visualizations cho paper (6 figures planned)
3. Verify IEMOCAP per-class F1 analysis""",
        "next_tasks": [
            ("Viết paper: Dataset + Experimental Setup", "Week 3 (May 26)"),
            ("Tạo 6 figures: learning curves, weight distribution, confusion matrix", "Week 3 (May 28)"),
            ("IEMOCAP per-class F1 + confusion matrix analysis", "Week 3 (May 27)"),
        ],
    },
    {
        "name": "Hồ Gia Phú",
        "id": "22110060",
        "filename": r"d:\OJT\Worklog\Worklog - W2 - 22110060 - HoGiaPhu.docx",
        "tasks": [
            ("Chạy ECR Ablation Study: 4 variants (full, no_certainty, ce_pseudo, no_augment) × 2 datasets × 3 seeds = 24 experiments",
             "Completed", "3h",
             "KL > CE pseudo-label (-0.61% MELD). Augmentation cần thiết (-1.11% IEMOCAP)"),
            ("Fix IEMOCAP 50% weakness: grid search 54 experiments (6 λ_u × 3 ramp-up × 3 seeds)",
             "Completed", "5h",
             "Best config: λ_u=0.3, rp=20 → WF1=0.5912 (gap vs FM giảm từ -0.8% xuống -0.5%)"),
            ("Implement CoMPM architecture: Transformer encoder + speaker-aware memory tracking cho baseline comparison",
             "Completed", "3h",
             "models/erc/sota_baselines.py — CoMPMEncoder với 4 local epochs, lr=5e-4"),
            ("Implement SPCL architecture: Prototypical Contrastive Learning + curriculum strategy",
             "Completed", "3h",
             "SPCLDialogueRNN với contrastive loss ramp-up, lambda_cl schedule"),
            ("Prepare presentation slides: 10-section LaTeX beamer deck với speaker notes",
             "Completed", "2h",
             "Present26OJT_slides.tex — professional narrative, research positioning"),
            ("Hỗ trợ debug training pipeline: fix evaluate_sota() cho CE-based models, decouple từ Dirichlet pipeline",
             "Completed", "2h",
             "evaluate_sota() outputs logits → CE loss, separate from EDL uncertainty pipeline"),
        ],
        "total_time": "~18h",
        "summary_tasks": """-Key tasks done:
1. Hoàn thành ECR Ablation (24 experiments): KL tốt hơn CE pseudo-label, strong augmentation cần thiết
2. Grid search IEMOCAP 50% fix (54 experiments): λ_u=0.3 + ramp_up=20 giảm gap vs FixMatch xuống -0.5%
3. Implement 2 SOTA baseline architectures: CoMPM (Transformer memory) + SPCL (Prototypical CL)
4. Prepare presentation slides với professional narrative
5. Fix evaluate_sota() để decouple CE models từ Dirichlet pipeline
6. Tổng cộng 78 experiments mới chạy thành công""",
        "summary_learned": """-Key things learned:
1. ECR Ablation: KL-based consistency > CE pseudo-labeling, certainty weighting có effect nhỏ nhưng mixed
2. IEMOCAP 50% insight: λ_u thấp (0.3) tốt hơn cao (1.0+) — SSL weight quá mạnh phá model
3. CoMPM architecture: Transformer self-attention + positional encoding cho speaker context tracking
4. SPCL: EMA-updated class prototypes + contrastive loss cần ramp-up schedule để stable
5. Presentation craft: research narrative quan trọng — đừng tự hạ bệ dự án trong slides""",
        "summary_literature": """-Literature read:
1. Lee & Lee (2022) "CoMPM" — Context Modeling with Speaker's Pre-trained Memory Tracking (NAACL)
2. Song et al. (2022) "SPCL" — Supervised Prototypical Contrastive Learning for ERC (EMNLP)
3. Shen et al. (2021) "DAG-ERC" — Directed Acyclic Graph for ERC (ACL)
4. Zhu et al. (2021) "TODKAT" — Topic-Driven Knowledge-Aware Transformer (ACL)
5. Han et al. (2021) "Trusted Multi-View Classification" — Dempster-Shafer theory (ICLR)""",
        "summary_issues": """-Issues/Challenges:
1. ECR no_certainty trên MELD cao hơn full ECR (+0.39%) — certainty weighting không consistent
2. IEMOCAP 50%: best ECR config (0.5912) vẫn thua FixMatch (0.5965) dù gap nhỏ
3. p-values ablation không significant (n=3 quá ít) — cần nâng lên 5 seeds
4. CoMPM/SPCL ban đầu shared weights với FedAvg → kết quả giả, đã fix
5. Presentation quá dài (17 slides) → cắt xuống 10 slides theo yêu cầu""",
        "next_plan": """Plan for the next week:
1. Viết paper: Methodology section (DialogueRNN + EDL + ECR formulation)
2. IEMOCAP 4-class experiment cho fair comparison với SOTA
3. Hỗ trợ tạo figures cho paper""",
        "next_tasks": [
            ("Viết paper: Method section (ECR + EAFA mathematical formulation)", "Week 3 (May 26)"),
            ("Chạy IEMOCAP 4-class experiment (merge happy+excited)", "Week 3 (May 27)"),
            ("Tạo confusion matrix + per-class F1 figures", "Week 3 (May 28)"),
        ],
    },
]


def generate_worklog(member):
    doc = docx.Document(TEMPLATE)

    # Table 0: Project Info
    info_table = doc.tables[0]
    info_table.rows[0].cells[1].text = PROJECT_NAME
    info_table.rows[1].cells[1].text = member["name"]
    info_table.rows[2].cells[1].text = member["id"]

    # Table 1: Weekly Tasks
    task_table = doc.tables[1]

    # Delete data rows (keep header rows 0,1,2)
    while len(task_table.rows) > 3:
        tr = task_table.rows[3]._tr
        task_table._tbl.remove(tr)

    # Add task rows
    for task, status, time_spent, note in member["tasks"]:
        row = task_table.add_row()
        row.cells[0].text = task
        row.cells[1].text = status
        row.cells[2].text = time_spent
        row.cells[3].text = note

    # Total row
    total_row = task_table.add_row()
    total_row.cells[0].text = "TOTAL WEEKLY TIME SPENT"
    total_row.cells[1].text = ""
    total_row.cells[2].text = member["total_time"]
    total_row.cells[3].text = ""

    # Dates — Week 2
    task_table.rows[0].cells[1].text = "2"
    task_table.rows[0].cells[2].text = "2"
    task_table.rows[0].cells[3].text = "2"
    task_table.rows[1].cells[1].text = "19/05/2026 - 23/05/2026"
    task_table.rows[1].cells[2].text = ""
    task_table.rows[1].cells[3].text = ""

    # Paragraphs
    doc.paragraphs[6].text = member["summary_tasks"]
    doc.paragraphs[8].text = member["summary_learned"]
    doc.paragraphs[10].text = member["summary_literature"]
    doc.paragraphs[12].text = member["summary_issues"]
    doc.paragraphs[15].text = member["next_plan"]

    # Table 2: Next Week Plan
    plan_table = doc.tables[2]
    while len(plan_table.rows) > 1:
        tr = plan_table.rows[1]._tr
        plan_table._tbl.remove(tr)

    for task, expected in member["next_tasks"]:
        row = plan_table.add_row()
        row.cells[0].text = task
        row.cells[1].text = expected

    doc.save(member["filename"])
    print(f"  Saved: {member['filename']}")


if __name__ == "__main__":
    print("Generating Work Logs for Week 2...")
    for m in members:
        generate_worklog(m)
    print("\nDone! 3 worklogs generated.")
