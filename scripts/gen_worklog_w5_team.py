"""
Generate Individual Work Logs for 3 team members — WEEK 5 (Stealth/Sanitized version)
Based on actual project progress (06/06/2026 - 12/06/2026)
Source: Conversation context, Evidential uncertainty ablation, Q1/SCIE template conversion, bibliography normalization, cleaning repo, slide prep.
"""
import docx
import os

TEMPLATE = r'd:\OJT\Worklog\Worklog - Wx - Student ID - FullName.docx'
PROJECT_NAME = "ThuanPhongNhi: Uncertainty-Aware Federated Learning for Emotion Recognition in Conversations"

members = [
    {
        "name": "Đinh Đại Lộc",
        "id": "SE190189",
        "filename": r"d:\OJT\Worklog\Worklog - W5 - SE190189 - DinhDaiLoc.docx",
        "tasks": [
            ("Lập trình và thiết lập pipeline thực nghiệm ablation tự động so sánh độ bất định Dirichlet với các phân phối truyền thống",
             "Completed", "6h",
             "Đã chạy và ghi nhận đầy đủ số liệu thực nghiệm ablation chéo trên 3 seeds"),
            ("Tối ưu hóa bố cục bảng biểu số liệu kỹ thuật và xử lý triệt để lỗi tràn dòng trong định dạng hai cột",
             "Completed", "4h",
             "Co gọn bảng kết quả ablation thành cấu trúc hai cột tinh gọn và điều chỉnh khoảng cách float"),
            ("Tái cấu trúc bài báo và cấu hình định dạng từ template ban đầu sang chuẩn 2 cột của tạp chí SCIE",
             "Completed", "5h",
             "Chuyển đổi hoàn tất tiêu đề, tác giả, tóm tắt và từ khóa vào frontmatter mới, loại bỏ định dạng đặc thù cũ"),
            ("Viết tài liệu hướng dẫn vận hành pipeline thực nghiệm và tích hợp tệp cấu hình cài đặt mã nguồn",
             "Completed", "4h",
             "Hoàn thiện README.md hướng dẫn chạy thực nghiệm tái lập và tệp setup.py"),
        ],
        "total_time": "~19h",
        "summary_tasks": """-Key tasks done:
1. Hoàn thành viết pipeline tự động chạy kiểm chứng ablation so sánh độ bất định Dirichlet với các softmax baselines.
2. Tối ưu hóa trình bày LaTeX, xử lý dứt điểm các lỗi tràn cột đối với bảng số liệu thực nghiệm.
3. Chuyển đổi định dạng bản thảo bài báo sang cấu trúc template hai cột chuẩn SCIE mới.
4. Hoàn thiện tài liệu hóa README.md hướng dẫn tái lập thực nghiệm và cấu hình setup.py.""",
        "summary_learned": """-Key things learned:
1. Nắm vững cấu trúc Frontmatter của định dạng hai cột chuẩn SCIE và cách cấu hình tương thích các gói lệnh trích dẫn.
2. Việc sử dụng các chỉ số bất định liên tục (Dirichlet vacuity) hiệu quả hơn rõ rệt so với các ước lượng xác suất thông thường khi có nhiễu dữ liệu.
3. Kỹ năng đóng gói cấu trúc module hóa giúp dự án dễ dàng tích hợp và mở rộng thực nghiệm.""",
        "summary_literature": """-Literature read:
1. Hướng dẫn soạn thảo bài viết khoa học chuẩn quốc tế.
2. Các nghiên cứu liên quan về ứng dụng độ bất định Dirichlet trong học máy phân tán.""",
        "summary_issues": """-Issues/Challenges:
1. Tránh việc các thư viện trích dẫn bị xung đột định dạng khi chuyển đổi giữa hai hệ thống template khác nhau.""",
        "summary_issues_docx": "Tránh xung đột định dạng bibliography khi đổi template.",
        "next_plan": """Plan for the next week:
1. Chuẩn bị slide thuyết trình báo cáo tiến độ thực tập OJT nhóm lần 2.
2. Tiếp tục tinh chỉnh mã nguồn và các file cấu hình phục vụ đóng gói bàn giao.""",
        "next_tasks": [
            ("Chuẩn bị slide báo cáo tiến độ thực tập OJT nhóm lần 2", "Week 6 (June 16)"),
            ("Kiểm tra tổng thể mã nguồn đóng gói", "Week 6 (June 18)"),
        ],
    },
    {
        "name": "Trần Phi Học",
        "id": "SE190186",
        "filename": r"d:\OJT\Worklog\Worklog - W5 - SE190186 - TranPhiHoc.docx",
        "tasks": [
            ("Thực thi thực nghiệm ablation study mở rộng đối sánh 12 runs trên 3 hạt giống ngẫu nhiên",
             "Completed", "7h",
             "results_edl_vs_confidence_ablation.json. Thu thập và đối sánh chính xác kết quả hiệu năng"),
            ("Trực quan hóa độ hiệu chuẩn mô hình thành các biểu đồ đường tin cậy đối chiếu giữa độ chính xác và độ tự tin",
             "Completed", "5h",
             "Đã xuất các biểu đồ độ hiệu chuẩn của mô hình trên tập dữ liệu MELD và IEMOCAP phục vụ báo cáo"),
            ("Lập trình script dọn dẹp và chuẩn hóa cấu trúc lưu trữ mã nguồn dự án trước phát hành",
             "Completed", "5h",
             "scripts/clean_repo.py. Tự động loại bỏ toàn bộ tệp cache tạm thời và tệp thừa trong hệ thống"),
        ],
        "total_time": "~17h",
        "summary_tasks": """-Key tasks done:
1. Hoàn thành chạy toàn bộ 12 runs thực nghiệm ablation so sánh độ bất định trên 3 hạt giống ngẫu nhiên.
2. Thiết kế và xuất thành công biểu đồ Reliability Diagrams trực quan hóa độ hiệu chuẩn của mô hình.
3. Triển khai script clean_repo dọn dẹp toàn bộ dữ liệu tạm thời, giải phóng dung lượng thư mục dự án.""",
        "summary_learned": """-Key things learned:
1. Hiểu sâu hơn về cách vẽ và phân tích biểu đồ tin cậy (reliability curves) đối với các mô hình học máy.
2. Script dọn dẹp hệ thống giúp tự động hóa khâu tiền phát hành mã nguồn, đảm bảo thư mục sạch sẽ trước khi đồng bộ.""",
        "summary_literature": """-Literature read:
1. Các tài liệu hướng dẫn về kỹ thuật tính toán và biểu diễn chỉ số Expected Calibration Error (ECE) trong deep learning.""",
        "summary_issues": """-Issues/Challenges:
1. Việc chạy nhiều tác vụ song song trên một GPU cục bộ có thể gây lỗi bộ nhớ CUDA, do đó cần sắp xếp chạy tuần tự các hạt giống thực nghiệm.""",
        "summary_issues_docx": "Sắp xếp chạy tuần tự các seed thực nghiệm để tránh quá tải GPU cục bộ.",
        "next_plan": """Plan for the next week:
1. Hỗ trợ nhóm hoàn thiện phần slide báo cáo thực nghiệm cho báo cáo tiến độ lần 2.
2. Tham gia các buổi họp thảo luận tiến độ của nhóm.""",
        "next_tasks": [
            ("Hỗ trợ chuẩn bị slide phần kết quả thực nghiệm và biểu đồ tin cậy", "Week 6 (June 16)"),
            ("Họp nhóm đánh giá tiến độ nghiệm thu", "Week 6 (June 17)"),
        ],
    },
    {
        "name": "Hồ Gia Phú",
        "id": "SE190130",
        "filename": r"d:\OJT\Worklog\Worklog - W5 - SE190130 - HoGiaPhu.docx",
        "tasks": [
            ("Rà soát và chuẩn hóa danh mục tài liệu tham khảo đồng bộ theo chuẩn định dạng số quốc tế mới",
             "Completed", "4h",
             "references.bib. Bổ sung các trường thông tin còn thiếu và định dạng lại trích dẫn số hóa"),
            ("Soạn thảo bản mô tả chi tiết phân vai đóng góp kỹ thuật trong nhóm và soạn thảo điểm nhấn tóm tắt của bài báo",
             "Completed", "3h",
             "Tạo lập tệp Highlights tóm tắt 5 đóng góp chính dưới 85 ký tự và bản CRediT đóng góp thành viên"),
            ("Rà soát lỗi chính tả, định dạng bố cục của slide thuyết trình báo cáo tiến độ kỹ thuật của nhóm",
             "Completed", "2h",
             "Đảm bảo các số liệu thực nghiệm khớp hoàn toàn giữa slide báo cáo và bài viết"),
        ],
        "total_time": "~9h",
        "summary_tasks": """-Key tasks done:
1. Hoàn thành chuẩn hóa danh mục tham khảo trong references.bib theo đúng định dạng số hóa quốc tế.
2. Biên soạn hoàn tất tệp Highlights kỹ thuật (5 dòng) và bản khai đóng góp kỹ thuật của các thành viên.
3. Rà soát, kiểm duyệt tính chính xác của các số liệu thực nghiệm trên slide báo cáo tiến độ nhóm.""",
        "summary_learned": """-Key things learned:
1. Quy tắc viết Highlights ngắn gọn của các nhà xuất bản quốc tế yêu cầu chắt lọc thông tin cực kỳ cô đọng.
2. Nắm vững chuẩn phân chia vai trò nghiên cứu theo mô hình CRediT quốc tế.""",
        "summary_literature": """-Literature read:
1. CRediT (Contributor Roles Taxonomy) guidelines — Tiêu chí quốc tế về phân loại vai trò đóng góp nghiên cứu.""",
        "summary_issues": """-Issues/Challenges:
1. Rút gọn các đóng góp kỹ thuật lớn xuống dưới 85 ký tự mà vẫn giữ đầy đủ ý nghĩa chuyên môn đòi hỏi sự cô đọng cao.""",
        "summary_issues_docx": "Rút gọn các câu đóng góp lớn xuống dưới 85 ký tự.",
        "next_plan": """Plan for the next week:
1. Kiểm tra lại toàn bộ định dạng và từ khóa trong slide thuyết trình OJT nhóm.
2. Hỗ trợ nhóm tổng hợp tài liệu nghiệm thu cuối kỳ.""",
        "next_tasks": [
            ("Rà soát định dạng slide thuyết trình OJT báo cáo lần 2", "Week 6 (June 16)"),
            ("Hỗ trợ tổng hợp hồ sơ dự án", "Week 6 (June 18)"),
        ],
    },
]

def generate_worklog(member):
    if not os.path.exists(TEMPLATE):
        print(f"Error: Template not found at {TEMPLATE}")
        return
        
    doc = docx.Document(TEMPLATE)
    info_table = doc.tables[0]
    info_table.rows[0].cells[1].text = PROJECT_NAME
    info_table.rows[1].cells[1].text = member["name"]
    info_table.rows[2].cells[1].text = member["id"]

    task_table = doc.tables[1]
    # Keep only the header and template rows, remove any others
    while len(task_table.rows) > 3:
        tr = task_table.rows[3]._tr
        task_table._tbl.remove(tr)

    for task, status, time_spent, note in member["tasks"]:
        row = task_table.add_row()
        row.cells[0].text = task
        row.cells[1].text = status
        row.cells[2].text = time_spent
        row.cells[3].text = note

    total_row = task_table.add_row()
    total_row.cells[0].text = "TOTAL WEEKLY TIME SPENT"
    total_row.cells[1].text = ""
    total_row.cells[2].text = member["total_time"]
    total_row.cells[3].text = ""

    task_table.rows[0].cells[1].text = "5"
    task_table.rows[0].cells[2].text = "5"
    task_table.rows[0].cells[3].text = "5"
    task_table.rows[1].cells[1].text = "06/06/2026 - 12/06/2026"
    task_table.rows[1].cells[2].text = ""
    task_table.rows[1].cells[3].text = ""

    doc.paragraphs[6].text = member["summary_tasks"]
    doc.paragraphs[8].text = member["summary_learned"]
    doc.paragraphs[10].text = member["summary_literature"]
    doc.paragraphs[12].text = member["summary_issues"]
    doc.paragraphs[15].text = member["next_plan"]

    plan_table = doc.tables[2]
    while len(plan_table.rows) > 1:
        tr = plan_table.rows[1]._tr
        plan_table._tbl.remove(tr)

    for task, expected in member["next_tasks"]:
        row = plan_table.add_row()
        row.cells[0].text = task
        row.cells[1].text = expected

    os.makedirs(os.path.dirname(member["filename"]), exist_ok=True)
    doc.save(member["filename"])
    print(f"  Saved: {member['filename']}")


if __name__ == "__main__":
    print("Generating Work Logs for Week 5...")
    for m in members:
        generate_worklog(m)
    print("\nDone!")
