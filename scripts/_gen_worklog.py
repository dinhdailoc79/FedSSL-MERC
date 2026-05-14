"""
Generate Week 1 Worklog - Fixed version
Rebuilds Table 1 properly to avoid row ordering issues
"""
import docx
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from copy import deepcopy

# Load template
doc = docx.Document(r'd:\OJT\Worklog - Wx - Student ID - FullName.docx')

# ============================================
# Table 0: Project Info
# ============================================
info_table = doc.tables[0]
info_table.rows[0].cells[1].text = "LucBinh: Uncertainty-Aware Federated Semi-Supervised Learning for Multimodal Emotion Recognition in Conversations"
info_table.rows[1].cells[1].text = "Dinh Dai Loc"
info_table.rows[2].cells[1].text = "22110046"

# ============================================
# Table 1: Weekly Tasks - Rebuild properly
# ============================================
task_table = doc.tables[1]

# Remove all existing rows except header rows (0,1,2) and total row
# Strategy: clear all content and rebuild
all_tasks = [
    ("Download & verify MELD dataset (1,432 dialogues, 13,708 utterances, 7 emotions) and IEMOCAP dataset (151 sessions, 5,531 utterances, 4 emotions)", 
     "Completed", "3h",
     "data/raw/MELD/, data/raw/IEMOCAP/ verified"),
    
    ("Extract RoBERTa-base text features for MELD (9,989+1,109+2,610 utts) and IEMOCAP (5,531 utts)",
     "Completed", "2h",
     "meld_text_roberta.pt (43.4MB), iemocap_text_roberta_4class.pt (17.8MB)"),
    
    ("Train centralized DialogueRNN baseline on MELD (50 epochs, early stopping)",
     "Completed", "2h",
     "Test WF1=0.5442, epoch 27/50. Checkpoint saved."),
    
    ("Implement FL infrastructure: FedAvg, FedProx, client, server",
     "Completed", "5h",
     "federated/aggregation/fedavg.py, fedprox.py, client.py, server.py"),
    
    ("Run FedAvg (5 clients, alpha=0.5, 30 rounds) and FedProx (mu=0.01)",
     "Completed", "1h",
     "FedAvg WF1=0.5419, FedProx WF1=0.5412 (~0.5% gap vs centralized)"),
    
    ("Implement FixMatch SSL: augmentation, curriculum threshold, train script",
     "Completed", "4h",
     "semi_supervised/augmentation.py, fixmatch.py, scripts/train_ssl.py"),

    ("Run SSL experiments at 10% and 20% label ratios",
     "Completed", "1h",
     "20% SSL WF1=0.4141 vs Supervised WF1=0.4736. Validates Poison Amplification."),
    
    ("Review proposal LucBinh, gap analysis vs implementation",
     "Completed", "2h",
     "Baselines done. Next: EDL module (core contribution)."),
]

# Delete all data rows (keep header rows 0,1,2 and we'll re-add total)
while len(task_table.rows) > 3:
    tr = task_table.rows[3]._tr
    task_table._tbl.remove(tr)

# Add task rows
for task, status, time_spent, note in all_tasks:
    row = task_table.add_row()
    row.cells[0].text = task
    row.cells[1].text = status
    row.cells[2].text = time_spent
    row.cells[3].text = note

# Add total row at end
total_row = task_table.add_row()
total_row.cells[0].text = "TOTAL WEEKLY TIME SPENT"
total_row.cells[1].text = ""
total_row.cells[2].text = "~20h"
total_row.cells[3].text = ""

# ============================================
# Dates
# ============================================
task_table.rows[0].cells[1].text = "1"
task_table.rows[0].cells[2].text = "1"
task_table.rows[0].cells[3].text = "1"
task_table.rows[1].cells[1].text = "09/05/2026 - 15/05/2026"
task_table.rows[1].cells[2].text = ""
task_table.rows[1].cells[3].text = ""

# ============================================
# Paragraphs: Summary/Reflection  
# ============================================
doc.paragraphs[6].text = """-Key tasks done:
1. Completed data pipeline: Downloaded, verified, and extracted RoBERTa features for both MELD and IEMOCAP datasets
2. Established centralized baseline: DialogueRNN achieves WF1=0.5442 on MELD test set
3. Implemented full Federated Learning infrastructure (FedAvg + FedProx) with Dirichlet non-IID partitioning
4. Validated FL performance: FedAvg WF1=0.5419, FedProx WF1=0.5412 (only 0.4-0.6% gap vs centralized)
5. Implemented FixMatch SSL with feature-level augmentation and curriculum threshold
6. All code committed and pushed to GitHub repository"""

doc.paragraphs[8].text = """-Key things learned:
1. Federated Learning simulation: How to simulate distributed training on a single GPU using client/server architecture with parameter aggregation
2. Non-IID data challenges: Dirichlet partitioning (alpha=0.5) creates heterogeneous data distributions that are realistic for real-world scenarios
3. Semi-supervised learning challenges in ERC: FixMatch's softmax-based pseudo-labeling struggles with dialogue-level emotion recognition due to confirmation bias
4. Feature extraction pipeline: Pre-extracting RoBERTa embeddings (768-dim) saves significant computation during training iterations
5. GPU memory management: Batch size tuning and gradient clipping are essential for stable training on consumer GPUs (RTX 4050, 6GB VRAM)"""

doc.paragraphs[10].text = """-Literature read and key things learned:
1. McMahan et al. (2017) "Communication-Efficient Learning of Deep Networks from Decentralized Data" - FedAvg algorithm
2. Li et al. (2020) "Federated Optimization in Heterogeneous Networks" - FedProx proximal term for non-IID
3. Sohn et al. (2020) "FixMatch: Simplifying Semi-Supervised Learning" - Confidence thresholding insufficient for ERC
4. Sensoy et al. (2018) "Evidential Deep Learning to Quantify Classification Uncertainty" - Dirichlet-based uncertainty
5. Han et al. (2021) "Trusted Multi-View Classification" - Dempster-Shafer combination for multi-modal fusion
6. Ghosal et al. (2019) "DialogueGCN" and Majumder et al. (2019) "DialogueRNN" - Conversational context modeling"""

doc.paragraphs[12].text = """-Issues/problems/Challenges:
1. UnicodeEncodeError on Windows console - Fixed by converting special chars to ASCII
2. Python built-in 'ssl' module name conflicts with our package - Renamed to 'semi_supervised/'
3. FixMatch produces 0 pseudo-labels initially (threshold too high) - Added curriculum warmup but noise remains
4. IEMOCAP loader needed fixes: consensus annotation parsing + transcription file loading
5. Key finding: FixMatch at 20% labels (WF1=0.4141) underperforms supervised-only (WF1=0.4736) - validates Poison Amplification hypothesis"""

doc.paragraphs[15].text = """Plan for the next week:
1. Implement Evidential Deep Learning (EDL) module: Dirichlet head (softplus -> evidence -> belief + uncertainty)
2. Implement Evidential Loss: Type-II MLE + KL regularization
3. Implement Evidential Consistency Regularization (ECR): certainty-weighted Dirichlet KL
4. Implement EAFA aggregation: uncertainty-aware server weights
5. Extract WavLM audio features for multimodal experiments
6. Begin Dempster-Shafer fusion module"""

# ============================================
# Table 2: Next Week Plan
# ============================================
plan_table = doc.tables[2]

# Clear existing
while len(plan_table.rows) > 1:
    tr = plan_table.rows[1]._tr
    plan_table._tbl.remove(tr)

next_tasks = [
    ("Implement EDL module (Dirichlet head + evidential loss)", "Week 2 (May 22)"),
    ("Implement ECR loss (certainty-weighted Dirichlet KL)", "Week 2 (May 22)"),
    ("Implement EAFA aggregation (uncertainty-aware)", "Week 2 (May 20)"),
    ("Extract WavLM audio features for MELD", "Week 2 (May 19)"),
    ("Begin DS fusion module for multimodal evidence", "Week 2-3 (May 25)"),
]

for task, expected in next_tasks:
    row = plan_table.add_row()
    row.cells[0].text = task
    row.cells[1].text = expected

# Save
output_path = r'd:\OJT\Worklog - W1 - 22110046 - DinhDaiLoc.docx'
doc.save(output_path)
print(f"Saved to: {output_path}")
