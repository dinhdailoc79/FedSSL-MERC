"""
Generate Individual Work Logs for 3 team members — WEEK 4
Based on actual project progress (31/05/2026 - 05/06/2026)
Source: Conversation context, AAAI Rebuttal preparation, ECE/NLL calibration, Byzantine robustness, FL baselines (SCAFFOLD/FedNova/FedAdam/MOON), significance tests.
"""
import docx
import os

TEMPLATE = r'd:\OJT\Worklog\Worklog - Wx - Student ID - FullName.docx'
PROJECT_NAME = "ThuanPhongNhi: Uncertainty-Aware Federated Learning for Emotion Recognition in Conversations"

members = [
    {
        "name": "Đinh Đại Lộc",
        "id": "SE190189",
        "filename": r"d:\OJT\Worklog\Worklog - W4 - SE190189 - DinhDaiLoc.docx",
        "tasks": [
            ("Rà soát, hoàn thiện và tích hợp toàn bộ chứng minh toán học (Appendix Proofs) cho Theorem 1 và Corollary 1 vào bản thảo fedssl-merc-advisor.tex",
             "Completed", "6h",
             "Đã gộp thành công phần phụ lục toán học vào cuối bài để gửi Thầy/Cô duyệt toàn bộ trong 1 PDF"),
            ("Giải quyết các lỗi ngắt từ và căn chỉnh bố cục các bảng (Table 4, 5, 6, 7, 8) trên LaTeX",
             "Completed", "5h",
             "Sử dụng mbox ngăn ngắt đôi từ 'quantifies' và 'scalability' qua trang, co gọn đoạn văn Controlled Centralized Comparisons để giải quyết triệt để lỗi tràn cột"),
            ("Soạn thảo bản nháp phản hồi phản biện (Author Response / Rebuttal) giải trình các ý kiến đóng góp của Reviewer",
             "Completed", "5h",
             "Giải quyết các câu hỏi cốt lõi về độ ổn định dưới low-label, cơ chế phòng vệ Byzantine không cần server validation set, và độ nhạy của ECR"),
            ("Cấu hình tệp .gitignore và dọn dẹp Git cache để untrack các thư mục chứa bài viết mật (docs/ và paper/)",
             "Completed", "3h",
             "Đảm bảo an toàn tuyệt đối cho các bản thảo mật và tài liệu của dự án khi đẩy code thực nghiệm lên GitHub công khai"),
        ],
        "total_time": "~19h",
        "summary_tasks": """-Key tasks done:
1. Hoàn thành tích hợp phần chứng minh toán học (Appendix Proofs) vào bản thảo advisor.tex.
2. Khắc phục hoàn toàn lỗi ngắt từ và lỗi trình bày bảng biểu trên LaTeX (Table 4, 5, 6, 7, 8).
3. Soạn thảo dự thảo Rebuttal giải trình chi tiết các thắc mắc của Reviewer về độ ổn định và an toàn của mô hình.
4. Cấu hình .gitignore và dọn dẹp Git cache giúp bảo mật tài liệu bài báo khi đồng bộ mã nguồn lên GitHub.""",
        "summary_learned": """-Key things learned:
1. Việc co gọn một vài từ không quan trọng trong đoạn văn (typesetting optimization) là kỹ năng thực tế rất hữu ích để giải quyết lỗi tràn dòng và cải thiện thẩm mỹ văn bản LaTeX.
2. Để tránh việc các thư mục chứa tài liệu mật bị đẩy lên GitHub, cần phải dọn dẹp Git cache sau khi cập nhật tệp .gitignore để các quy tắc bỏ qua có hiệu lực ngay lập tức.
3. Bản review phản hồi "Weak Accept" chỉ ra các điểm yếu về kiểm định thống kê và bảo mật mà nhóm cần đặc biệt lưu tâm.""",
        "summary_literature": """-Literature read:
1. AAAI 2026 Author Kit & Formatting Guidelines — Quy chuẩn định dạng văn bản và giới hạn trang của hội nghị.
2. Peer Review comments — Phân tích các tiêu chí đánh giá của phản biện học thuật đối với mô hình FedSSL-MERC.""",
        "summary_issues": """-Issues/Challenges:
1. Việc giảm kích thước minipage của bảng làm tiêu đề tự động xuống dòng nhiều hơn, dẫn đến tăng chiều cao của bảng và đẩy chữ sang cột khác.
2. Các thư mục bị Git theo dõi từ trước sẽ không tự động bị loại bỏ chỉ bằng cách thêm vào .gitignore nếu không chạy lệnh xóa cache.""",
        "next_plan": """Plan for the next week:
1. Chuẩn bị nội dung thuyết trình báo cáo OJT lần 2 dựa trên kết quả tối ưu hóa bài báo.
2. Tiếp tục hoàn thiện tài liệu hướng dẫn chạy code thực nghiệm.""",
        "next_tasks": [
            ("Chuẩn bị slide báo cáo tiến độ và kết quả tối ưu hóa bài báo lần 2", "Week 5 (June 9)"),
            ("Rà soát và làm sạch mã nguồn thực nghiệm trước khi gửi Thầy/Cô kiểm tra", "Week 5 (June 11)"),
        ],
    },
    {
        "name": "Trần Phi Học",
        "id": "SE190186",
        "filename": r"d:\OJT\Worklog\Worklog - W4 - SE190186 - TranPhiHoc.docx",
        "tasks": [
            ("Thiết lập và chạy thực nghiệm đánh giá độ tin cậy độ bất định (Calibration Analysis: ECE & NLL) trên MELD",
             "Completed", "7h",
             "results/calibration_summary.json. Evidential head đạt mức sai số hiệu chỉnh ECE cực thấp (5.06%) khi dùng features finetuned"),
            ("Chạy thực nghiệm Byzantine & Client Spoofing Robustness giả lập tấn công đảo ngược gradient và giả mạo độ bất định",
             "Completed", "5h",
             "results_byzantine_robustness.json. EAFA tự động suy hao trọng số của client tấn công từ 0.39 xuống 0.07"),
            ("Thực hiện thực nghiệm so sánh EAFA với các FL optimizers hiện đại (SCAFFOLD, FedNova, FedAdam, MOON) dưới 10% nhãn",
             "Completed", "5h",
             "results/fl_baselines_results.json. EAFA đạt hiệu năng cạnh tranh vượt trội trong môi trường heterogeneous/noise"),
        ],
        "total_time": "~17h",
        "summary_tasks": """-Key tasks done:
1. Hoàn thành thực nghiệm đo đạc độ hiệu chuẩn ECE và NLL để chứng minh tính tin cậy của độ bất định ước lượng từ evidential head.
2. Mô phỏng thành công kịch bản tấn công Byzantine (label-flipping và sign-flipping) để kiểm chứng khả năng bảo mật của EAFA.
3. Chạy thực nghiệm so sánh với 4 thuật toán tối ưu FL baselines mạnh trên MELD và IEMOCAP.""",
        "summary_learned": """-Key things learned:
1. Độ bất định ước lượng bằng Evidential Deep Learning tương ứng rất tốt với chất lượng dữ liệu của client, giúp hệ thống tự động suy hao trọng số của client nhiễu hoặc độc hại.
2. EAFA có lợi thế truyền thông cực lớn khi chỉ truyền thêm đúng 1 tham số scalar (độ bất định) mà vẫn đạt hiệu năng cạnh tranh với các phương pháp phức tạp như SCAFFOLD hay MOON.""",
        "summary_literature": """-Literature read:
1. Karimireddy et al. (2020) "SCAFFOLD: Stochastic Controlled Averaging for Federated Learning".
2. Wang et al. (2020) "Tackling Objective Inconsistency in Federated Optimization (FedNova)".
3. Li et al. (2021) "Model-Contrastive Federated Learning (MOON)".""",
        "summary_issues": """-Issues/Challenges:
1. Khi chạy thực nghiệm so sánh, một số baseline đòi hỏi cấu hình hyperparameter rất khắt khe (như hệ số proximal của FedProx hay learning rate phía server của FedAdam).
2. Tấn công giả mạo độ bất định (spoofing) đòi hỏi mô hình global có cơ chế kiểm chuẩn chéo hoặc đồng thuận thời gian để phát hiện bất thường.""",
        "next_plan": """Plan for the next week:
1. Trực quan hóa kết quả phân tích độ hiệu chuẩn thành biểu đồ độ tin cậy (reliability curves) cho slide thuyết trình.
2. Hỗ trợ nhóm thiết kế slide OJT phần thực nghiệm.""",
        "next_tasks": [
            ("Vẽ các biểu đồ độ tin cậy (reliability curves) cho MELD general vs finetuned features", "Week 5 (June 9)"),
            ("Hỗ trợ chuẩn bị slide và tài liệu thuyết trình báo cáo tiến độ lần 2", "Week 5 (June 9)"),
        ],
    },
    {
        "name": "Hồ Gia Phú",
        "id": "SE190130",
        "filename": r"d:\OJT\Worklog\Worklog - W4 - SE190130 - HoGiaPhu.docx",
        "tasks": [
            ("Thực hiện kiểm định ý nghĩa thống kê (Statistical Significance Testing: Paired t-test, Wilcoxon, Cohen's d) trên 5 seeds",
             "Completed", "5h",
             "results/significance_tests.json. Đạt p-value < 0.05 đối với ECR vs FixMatch trên MELD, khẳng định kết quả có ý nghĩa thống kê"),
            ("Tích hợp các bảng số liệu thực nghiệm bổ sung (Calibration, Byzantine Robustness, FL Baselines) vào tệp LaTeX chính thức",
             "Completed", "3h",
             "Đảm bảo các bảng số liệu Table 5, Table 6 hiển thị chính xác và đồng bộ theo đúng cấu trúc bài báo"),
        ],
        "total_time": "~8h",
        "summary_tasks": """-Key tasks done:
1. Hoàn thành chạy kiểm định thống kê và tính toán p-value, Cohen's d, khoảng tin cậy (CI) cho các kết quả so sánh ECR vs FixMatch và EAFA vs FedAvg.
2. Tích hợp và đồng bộ các bảng số liệu kết quả thực nghiệm mới vào tệp LaTeX.""",
        "summary_learned": """-Key things learned:
1. Các phép kiểm định phi tham số (như Wilcoxon signed-rank test) rất quan trọng khi số lượng mẫu seeds nhỏ (N=5) để tránh giả định phân phối chuẩn.
2. Việc sắp xếp thứ tự và minipage trong LaTeX quyết định cách phân bổ bảng biểu vào các cột văn bản để tối ưu không gian trình bày.""",
        "summary_literature": """-Literature read:
1. Demšar (2006) "Statistical Comparisons of Classifiers over Multiple Data Sets" — Hướng dẫn thực hành kiểm định thống kê trong Machine Learning.
2. Wilcoxon (1945) "Individual Comparisons by Ranking Methods" — Lý thuyết về kiểm định Wilcoxon.""",
        "summary_issues": """-Issues/Challenges:
1. Với số lượng seeds nhỏ (N=3 đối với DailyDialog), kiểm định Wilcoxon không thể tính được giá trị p-value chính xác do cỡ mẫu tối thiểu yêu cầu lớn hơn.""",
        "next_plan": """Plan for the next week:
1. Rà soát lại toàn bộ định dạng bibliography và citation trong tệp references.bib.
2. Hỗ trợ nhóm rà soát lại bố cục trang in ấn của slide báo cáo OJT.""",
        "next_tasks": [
            ("Rà soát và chuẩn hóa tệp references.bib theo đúng định dạng AAAI", "Week 5 (June 9)"),
            ("Kiểm tra định dạng và chính tả cho slide thuyết trình OJT nhóm", "Week 5 (June 9)"),
        ],
    },
]


def generate_worklog(member):
    doc = docx.Document(TEMPLATE)
    info_table = doc.tables[0]
    info_table.rows[0].cells[1].text = PROJECT_NAME
    info_table.rows[1].cells[1].text = member["name"]
    info_table.rows[2].cells[1].text = member["id"]

    task_table = doc.tables[1]
    # Keep only the header and template rows, remove any others
    while len(task_table.rows) > 3:
        tr = task_table.rows[3]._tr
        task_table._tbl.remove(tr)

    for task, status, time_spent, note in member["tasks"]:
        row = task_table.add_row()
        row.cells[0].text = task
        row.cells[1].text = status
        row.cells[2].text = time_spent
        row.cells[3].text = note

    total_row = task_table.add_row()
    total_row.cells[0].text = "TOTAL WEEKLY TIME SPENT"
    total_row.cells[1].text = ""
    total_row.cells[2].text = member["total_time"]
    total_row.cells[3].text = ""

    task_table.rows[0].cells[1].text = "4"
    task_table.rows[0].cells[2].text = "4"
    task_table.rows[0].cells[3].text = "4"
    task_table.rows[1].cells[1].text = "31/05/2026 - 05/06/2026"
    task_table.rows[1].cells[2].text = ""
    task_table.rows[1].cells[3].text = ""

    doc.paragraphs[6].text = member["summary_tasks"]
    doc.paragraphs[8].text = member["summary_learned"]
    doc.paragraphs[10].text = member["summary_literature"]
    doc.paragraphs[12].text = member["summary_issues"]
    doc.paragraphs[15].text = member["next_plan"]

    plan_table = doc.tables[2]
    while len(plan_table.rows) > 1:
        tr = plan_table.rows[1]._tr
        plan_table._tbl.remove(tr)

    for task, expected in member["next_tasks"]:
        row = plan_table.add_row()
        row.cells[0].text = task
        row.cells[1].text = expected

    doc.save(member["filename"])
    print(f"  Saved: {member['filename']}")


if __name__ == "__main__":
    print("Generating Work Logs for Week 4...")
    for m in members:
        generate_worklog(m)
    print("\nDone!")
