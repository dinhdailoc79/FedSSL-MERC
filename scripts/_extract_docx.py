import docx
doc = docx.Document(r'd:\OJT\Worklog - Wx - Student ID - FullName.docx')

with open(r'd:\OJT\FedSSL-MERC\scripts\_worklog_template.txt', 'w', encoding='utf-8') as f:
    # Extract all paragraphs
    f.write("=== PARAGRAPHS ===\n")
    for i, para in enumerate(doc.paragraphs):
        f.write(f"P{i}: [{para.style.name}] {para.text}\n")
    
    # Extract all tables
    f.write("\n=== TABLES ===\n")
    for t_idx, table in enumerate(doc.tables):
        f.write(f"\n--- Table {t_idx} ({len(table.rows)} rows x {len(table.columns)} cols) ---\n")
        for r_idx, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            f.write(f"  Row {r_idx}: {' | '.join(cells)}\n")

print("Done")
