"""
Generate Individual Work Logs for 3 team members — WEEK 3
Based on actual project progress (26/05/2026 - 30/05/2026)
Source: Conversation context, SSL multi-ratio, DS Fusion missing modality, client scalability.
"""
import docx
import os

TEMPLATE = r'd:\OJT\Worklog\Worklog - Wx - Student ID - FullName.docx'
PROJECT_NAME = "ThuanPhongNhi: Uncertainty-Aware Federated Learning for Emotion Recognition in Conversations"

members = [
    {
        "name": "Đinh Đại Lộc",
        "id": "SE190189",
        "alt_id": "22110046",
        "filename": r"d:\OJT\Worklog\Worklog - W3 - SE190189 - DinhDaiLoc.docx",
        "alt_filename": r"d:\OJT\Worklog\Worklog - W3 - 22110046 - DinhDaiLoc.docx",
        "tasks": [
            ("Thiết lập missing-modality test harness (run_fusion_robustness.py) đánh giá DS Fusion dưới tỷ lệ mất audio (0%, 20%, 40%, 60%, 80%, 100%)",
             "Completed", "6h",
             "54 experiments xong. DS Fusion chỉ giảm 0.1% F1 ở 100% missing audio, chứng minh khả năng kháng nhiễu vượt trội"),
            ("Fix lỗi hệ thống huấn luyện SSL: sửa constructor EvidentialDialogueRNN và SupervisedEvidentialLoss",
             "Completed", "4h",
             "Giải quyết triệt để lỗi runtime crash liên quan đến annealing_epochs và constructor parameters"),
            ("Thực hiện ECR Ablation Study (5 seeds) và tái cấu trúc reframe certainty weighting thành 'safety mechanism'",
             "Completed", "5h",
             "Chứng minh w/o certainty weighting sụt giảm nghiêm trọng (-2.98%) trên IEMOCAP do dữ liệu nhiễu/lệch phân phối"),
            ("Cập nhật kết quả thí nghiệm DS Fusion và ECR Ablation vào manuscript LaTeX (advisor và submission version)",
             "Completed", "4h",
             "Đồng bộ Table 7 (Fusion Robustness) và phân tích lý thuyết. Nâng tổng số experiments từ 733 lên 805"),
        ],
        "total_time": "~19h",
        "summary_tasks": """-Key tasks done:
1. Thiết lập test harness đánh giá missing modality (54 runs): DS Fusion vs Concat/Attention.
2. Khắc phục lỗi constructor EvidentialDialogueRNN và Evidential Loss giúp pipeline chạy trơn tru.
3. Chạy 5-seed ECR Ablation Study và định vị lại certainty weighting như một 'safety mechanism' bảo vệ model khi dữ liệu nhiễu.
4. Cập nhật và đồng bộ kết quả nghiên cứu vào bản thảo AAAI (advisor & submission tex files), nâng tổng experiments lên 805.""",
        "summary_learned": """-Key things learned:
1. DS Fusion kết hợp bằng chứng cấp độ evidence (Dirichlet alphas) tự động suy hao về text-only khi audio evidence = 0 mà không làm sụp đổ hệ thống.
2. Certainty weighting là cơ chế an toàn cực kỳ quan trọng cho SSL khi đối mặt với dữ liệu nhiễu (noise) hoặc imbalance, dù trên dữ liệu sạch nó có vẻ trung tính.
3. Việc quản lý codebase federated đòi hỏi tính đồng bộ rất cao giữa các module độc lập.""",
        "summary_literature": """-Literature read:
1. Han et al. (2021) "Trusted Multi-View Classification" — Cơ sở toán học cho Dempster-Shafer theory trong EDL.
2. Sensoy et al. (2018) "Evidential Deep Learning to Quantify Classification Uncertainty" — Nền tảng về Dirichlet distribution.
3. Poria et al. (2019) "MELD: A Multimodal EmotionLines Dataset in Conversations" — Dataset audio-visual-text.""",
        "summary_issues": """-Issues/Challenges:
1. Audio modality trong MELD đóng góp rất ít vào độ chính xác chung, dẫn đến hiệu ứng trần (ceiling) cho cả 3 baseline.
2. ECR không có certainty weighting trên MELD cho kết quả hơi cao hơn (+0.28%) do tập dữ liệu sạch, dễ gây ngộ nhận nếu không chạy thêm trên IEMOCAP để đối chứng.""",
        "next_plan": """Plan for the next week:
1. Rà soát lại các phần chứng minh toán học trong Appendix.
2. Chuẩn bị tài liệu kỹ thuật cho việc thuyết trình slide OJT.""",
        "next_tasks": [
            ("Rà soát và hoàn thiện Appendix Proofs cho Theorem 1", "Week 4 (June 2)"),
            ("Hỗ trợ team chuẩn bị nội dung thuyết trình báo cáo Week 4", "Week 4 (June 4)"),
        ],
    },
    {
        "name": "Trần Phi Học",
        "id": "SE190186",
        "alt_id": "22110024",
        "filename": r"d:\OJT\Worklog\Worklog - W3 - SE190186 - TranPhiHoc.docx",
        "alt_filename": r"d:\OJT\Worklog\Worklog - W3 - 22110024 - TranPhiHoc.docx",
        "tasks": [
            ("Thực hiện chiến dịch huấn luyện SSL quy mô lớn (SSL Ratios Campaign) trên cả 3 datasets với 3 tỉ lệ nhãn qua 3 seeds",
             "Completed", "8h",
             "81 experiments chạy thành công. Kết quả lưu tại results_ssl_ratios.json"),
            ("Phân tích kết quả thực nghiệm SSL và viết báo cáo statistical significance, so sánh ECR vs FixMatch và Supervised",
             "Completed", "4h",
             "Chứng minh ECR vượt trội FixMatch ở tỉ lệ nhãn thấp (5%) do FixMatch bị kẹt bởi ngưỡng pseudo-label cứng (0.95) trong FL non-IID"),
            ("Thiết lập controlled comparison cho IEMOCAP 4-class bằng cách đồng bộ hóa feature extractor (RoBERTa finetuned)",
             "Completed", "4h",
             "So sánh công bằng với baseline DialogueRNN dùng chung features, kiểm chứng hiệu quả vượt trội thực tế của EAFA"),
        ],
        "total_time": "~16h",
        "summary_tasks": """-Key tasks done:
1. Huấn luyện thành công 81 experiments của chiến dịch SSL Ratios Campaign trên 3 datasets.
2. Phân tích kết quả thực nghiệm, lập bảng so sánh và chứng minh tính vượt trội của ECR vs FixMatch.
3. Thiết lập controlled comparison cho IEMOCAP 4-class dùng chung RoBERTa finetuned features để đảm bảo tính công bằng khoa học.""",
        "summary_learned": """-Key things learned:
1. FixMatch bị confirmation bias rất nặng trong FL khi clients có phân phối non-IID, do ngưỡng pseudo-label cứng chặn đứng việc sinh nhãn giả ở các round đầu.
2. Việc đồng bộ hóa feature extractor (encoder size, finetuned state) là yếu tố quyết định để đảm bảo tính công bằng khi so sánh với các SOTA baselines.""",
        "summary_literature": """-Literature read:
1. Sohn et al. (2020) "FixMatch: Simplifying Semi-Supervised Learning with Consistency and Confidence".
2. Zhang et al. (2021) "FlexMatch: Boosting Semi-Supervised Learning with Curriculum Pseudo-Labeling".
3. Lee & Lee (2022) "CoMPM: Context Modeling with Speaker's Pre-trained Memory".""",
        "summary_issues": """-Issues/Challenges:
1. DailyDialog 5% label ratio cho kết quả rất sát nhau giữa các phương pháp do hiệu ứng hội tụ sớm của text features chất lượng cao.
2. FixMatch thất bại hoàn toàn trong việc học ở các round đầu của IEMOCAP khi nhãn cực kỳ khan hiếm.""",
        "next_plan": """Plan for the next week:
1. Vẽ các biểu đồ trực quan hóa kết quả SSL Ratios cho paper.
2. Hỗ trợ chuẩn bị slide thuyết trình tiến độ.""",
        "next_tasks": [
            ("Tạo biểu đồ trực quan hóa (learning curves) cho 81 SSL experiments", "Week 4 (June 2)"),
            ("Chuẩn bị slide báo cáo kết quả thực nghiệm SSL", "Week 4 (June 4)"),
        ],
    },
    {
        "name": "Hồ Gia Phú",
        "id": "22110060",
        "alt_id": "SE190130",  # Just in case they need an FPT style ID as well
        "filename": r"d:\OJT\Worklog\Worklog - W3 - 22110060 - HoGiaPhu.docx",
        "alt_filename": r"d:\OJT\Worklog\Worklog - W3 - SE190130 - HoGiaPhu.docx",
        "tasks": [
            ("Chạy thực nghiệm độ giãn nở client (Client Scalability) với K=5, K=10, K=20 trên MELD và IEMOCAP (18 experiments)",
             "Completed", "5h",
             "results_client_scalability.json. EAFA tăng trưởng nhẹ (+0.3% ở K=20) trên MELD nhờ tác dụng ensemble"),
            ("Thực hiện notation cleanup trong toàn bộ bản thảo LaTeX, chuẩn hóa ký hiệu w_k (uncertainty weight) vs p_k",
             "Completed", "4h",
             "Đảm bảo tính chặt chẽ về mặt toán học trong Methodology và Appendix Proofs"),
            ("Thảo luận và viết 2 paragraphs về 'Fairness & Privacy' bổ sung vào phần Discussion của paper",
             "Completed", "3h",
             "Phân tích chi tiết khía cạnh đạo đức, rủi ro rò rỉ gradient/uncertainty trong hệ thống Federated MERC"),
        ],
        "total_time": "~12h",
        "summary_tasks": """-Key tasks done:
1. Thực hiện thực nghiệm Client Scalability (K=5/10/20) kiểm chứng khả năng scale của thuật toán EAFA.
2. Rà soát và làm sạch ký hiệu toán học (notation cleanup) trong toàn bộ paper.
3. Viết nội dung thảo luận Fairness & Privacy bổ sung cho manuscript.""",
        "summary_learned": """-Key things learned:
1. K=20 là quá lớn đối với các bộ dữ liệu hội thoại nhỏ như IEMOCAP (chỉ có 151 dialogues), gây ra tình trạng thiếu hụt dữ liệu nghiêm trọng ở mỗi client (Dirichlet partition).
2. Sự rõ ràng trong ký hiệu toán học (notation) quyết định tính dễ đọc và độ tin cậy của phần chứng minh hội tụ (convergence analysis).""",
        "summary_literature": """-Literature read:
1. Li et al. (2021) "A Field Guide to Federated Optimization" — Nền tảng về FL client scalability.
2. McMahan et al. (2018) "Learning Private Models with Differential Privacy".
3. Feng et al. (2022) "Semi-FedSER: Federated Semi-Supervised SER".""",
        "summary_issues": """-Issues/Challenges:
1. Hiện tượng thiếu hụt dữ liệu cực đoan tại client khi chia nhỏ tập dữ liệu IEMOCAP cho K=20 clients làm sụp tiến trình huấn luyện do một số client không có đủ class.""",
        "next_plan": """Plan for the next week:
1. Rà soát lại phần định dạng tài liệu tham khảo (citations/bibliography) cho paper.
2. Cập nhật các bảng số liệu Client Scalability vào manuscript.""",
        "next_tasks": [
            ("Rà soát lỗi định dạng trích dẫn (citations) trong references.bib", "Week 4 (June 2)"),
            ("Tích hợp bảng Client Scalability vào bản thảo LaTeX chính thức", "Week 4 (June 3)"),
        ],
    },
]


def generate_worklog(member):
    # Save primary filename
    doc = docx.Document(TEMPLATE)
    info_table = doc.tables[0]
    info_table.rows[0].cells[1].text = PROJECT_NAME
    info_table.rows[1].cells[1].text = member["name"]
    info_table.rows[2].cells[1].text = member["id"]

    task_table = doc.tables[1]
    while len(task_table.rows) > 3:
        tr = task_table.rows[3]._tr
        task_table._tbl.remove(tr)

    for task, status, time_spent, note in member["tasks"]:
        row = task_table.add_row()
        row.cells[0].text = task
        row.cells[1].text = status
        row.cells[2].text = time_spent
        row.cells[3].text = note

    total_row = task_table.add_row()
    total_row.cells[0].text = "TOTAL WEEKLY TIME SPENT"
    total_row.cells[1].text = ""
    total_row.cells[2].text = member["total_time"]
    total_row.cells[3].text = ""

    task_table.rows[0].cells[1].text = "3"
    task_table.rows[0].cells[2].text = "3"
    task_table.rows[0].cells[3].text = "3"
    task_table.rows[1].cells[1].text = "26/05/2026 - 30/05/2026"
    task_table.rows[1].cells[2].text = ""
    task_table.rows[1].cells[3].text = ""

    doc.paragraphs[6].text = member["summary_tasks"]
    doc.paragraphs[8].text = member["summary_learned"]
    doc.paragraphs[10].text = member["summary_literature"]
    doc.paragraphs[12].text = member["summary_issues"]
    doc.paragraphs[15].text = member["next_plan"]

    plan_table = doc.tables[2]
    while len(plan_table.rows) > 1:
        tr = plan_table.rows[1]._tr
        plan_table._tbl.remove(tr)

    for task, expected in member["next_tasks"]:
        row = plan_table.add_row()
        row.cells[0].text = task
        row.cells[1].text = expected

    doc.save(member["filename"])
    print(f"  Saved: {member['filename']}")


if __name__ == "__main__":
    print("Generating Work Logs for Week 3...")
    for m in members:
        generate_worklog(m)
    print("\nDone!")
