"""
Generate Individual Work Logs for 3 team members
Based on actual project progress as of Week 1 (09/05/2026 - 15/05/2026)
"""
import docx

TEMPLATE = r'd:\OJT\Worklog - Wx - Student ID - FullName.docx'
PROJECT_NAME = "LucBinh: Uncertainty-Aware Federated Learning for Emotion Recognition in Conversations"

# ============================================================
# Phân chia công việc CÂN ĐỐI cho 3 người
# ============================================================

members = [
    {
        "name": "Đinh Đại Lộc",
        "id": "22110046",
        "filename": r"d:\OJT\Worklog - W1 - 22110046 - DinhDaiLoc.docx",
        "tasks": [
            ("Thiết lập repository GitHub, cấu trúc dự án (data/, models/, federated/, scripts/), CI/CD pipeline",
             "Completed", "3h",
             "https://github.com/dinhdailoc79/FedSSL-MERC — cấu trúc module hóa 4 tầng"),
            ("Implement module Evidential Deep Learning: EDL Head (Dirichlet prediction), EDL Loss + KL regularization",
             "Completed", "5h",
             "models/evidential/edl_head.py, losses.py — Softplus → evidence → belief + uncertainty"),
            ("Implement EAFA: Evidence-Aware Federated Aggregation — client uncertainty → aggregation weights",
             "Completed", "4h",
             "federated/aggregation/eafa.py — beta-weighted exponential uncertainty mapping"),
            ("Fine-tune RoBERTa trên Kaggle GPU cho MELD, IEMOCAP, DailyDialog",
             "Completed", "3h",
             "scripts/finetune_roberta.py — MELD WF1=0.6380, IEMOCAP WF1=0.5431, DD WF1=0.8774"),
            ("Chạy thí nghiệm 3 seeds (42, 123, 2024) trên 3 datasets, thu thập mean±std",
             "Completed", "4h",
             "EAFA: MELD 63.09±0.07, IEMOCAP 60.44±0.25, DailyDialog 88.69±0.18"),
            ("Thu thập và phân tích số liệu SOTA từ papers (COSMIC, EmoBERTa, DAG-ERC, DialogueRNN)",
             "Completed", "2h",
             "Bảng so sánh 6 baselines. EAFA đạt 96.8% SOTA centralized trên MELD"),
        ],
        "total_time": "~21h",
        "summary_tasks": """-Key tasks done:
1. Thiết lập toàn bộ codebase và GitHub repository với cấu trúc module hóa
2. Implement core contributions: EDL Head (uncertainty prediction) + EAFA (uncertainty-aware aggregation)
3. Fine-tune RoBERTa trên Kaggle T4 GPU cho cả 3 datasets (MELD, IEMOCAP, DailyDialog)
4. Chạy 3-seed evaluation (seed 42, 123, 2024) → kết quả stable (std < 0.5%)
5. EAFA outperforms EDL centralized trên cả 3 datasets: +0.35% MELD, +2.13% IEMOCAP, +0.70% DD
6. Thu thập SOTA baselines và xây dựng comparison table""",
        "summary_learned": """-Key things learned:
1. Evidential Deep Learning: Dirichlet distribution cho phép đo uncertainty từ model predictions, khác biệt với softmax
2. Federated aggregation strategies: EAFA dùng uncertainty để weight clients tốt hơn FedAvg trung bình đơn giản
3. Kaggle workflow: Chia pipeline thành fine-tune (GPU) + train (CPU) để tối ưu resources
4. Evaluation rigor: 3-seed mean±std là yêu cầu bắt buộc cho top-tier venues (AAAI, IJCAI)
5. Feature extraction: Fine-tuned RoBERTa features tăng 3-7% WF1 so với frozen embeddings""",
        "summary_literature": """-Literature read:
1. Sensoy et al. (2018) "Evidential Deep Learning to Quantify Classification Uncertainty" — Dirichlet-based uncertainty
2. McMahan et al. (2017) "Communication-Efficient Learning" — FedAvg algorithm foundation
3. Ghosal et al. (2020) "COSMIC: COMmonSense knowledge for eMotion Identification" — SOTA baseline 65.21% MELD
4. Kim & Vossen (2021) "EmoBERTa" — RoBERTa fine-tuning for ERC, baseline 65.61% MELD
5. Majumder et al. (2019) "DialogueRNN" — Backbone model architecture""",
        "summary_issues": """-Issues/Challenges:
1. IEMOCAP fine-tune trên local bị lỗi data loader → giải quyết bằng export CSV + chạy trên Kaggle
2. Feature ID mismatch giữa fine-tuned output và pipeline → tạo fix_finetuned_features.py
3. IEMOCAP performance (60.44%) thấp hơn MELD do dataset nhỏ và 6-class khó hơn
4. DailyDialog evaluation protocol khác papers (weighted F1 vs micro F1 excl. neutral)
5. SSL (FixMatch) code tồn tại nhưng chưa tích hợp — cần quyết định bỏ/implement""",
        "next_plan": """Plan for the next week:
1. Chạy ablation study: FedAvg baseline + Softmax baseline để isolate contribution
2. Bắt đầu viết paper LaTeX (Introduction, Methodology)
3. Tối ưu hyperparameters cho IEMOCAP (learning rate, epochs)""",
        "next_tasks": [
            ("Chạy ablation study (FedAvg + Softmax baselines)", "Week 2 (May 20)"),
            ("Viết paper LaTeX: Introduction + Related Work", "Week 2 (May 22)"),
            ("Tối ưu IEMOCAP hyperparameters", "Week 2 (May 19)"),
        ],
    },
    {
        "name": "Trần Phi Học",
        "id": "22110024",
        "filename": r"d:\OJT\Worklog - W1 - 22110024 - TranPhiHoc.docx",
        "tasks": [
            ("Download, xác minh và tiền xử lý MELD dataset (1,432 dialogues, 13,708 utterances, 7 emotions)",
             "Completed", "2h",
             "data/raw/MELD/ — train/dev/test CSV verified, 7 emotion labels balanced check"),
            ("Download, xác minh IEMOCAP dataset (5 sessions, 151 dialogues, 7,380 utterances, 6 emotions)",
             "Completed", "3h",
             "data/raw/IEMOCAP/ — session 1-5, consensus annotations, 6-class emotion mapping"),
            ("Download, xác minh DailyDialog dataset (13,118 dialogues, 102,979 utterances, 7 emotions)",
             "Completed", "2h",
             "data/raw/DailyDialog/ — train/validation/test CSV, exclude no_emotion dialogues"),
            ("Implement data loaders cho 3 datasets: MELDDataset, IEMOCAPDataset, DailyDialogDataset",
             "Completed", "5h",
             "data/datasets/meld.py, iemocap.py, dailydialog.py — unified Dialogue/Utterance interface"),
            ("Implement Federated Data Partition: Dirichlet non-IID splitting cho N clients",
             "Completed", "3h",
             "data/federated_partition.py — alpha=0.5 Dirichlet, 5 clients, verified class distribution"),
            ("Extract frozen RoBERTa-base text features cho 3 datasets (768-dim embeddings)",
             "Completed", "3h",
             "meld_text_roberta.pt (43MB), iemocap_text_roberta.pt (23MB), dailydialog_text_roberta.pt (311MB)"),
            ("Export IEMOCAP sang CSV format cho Kaggle fine-tuning workflow",
             "Completed", "1h",
             "scripts/export_iemocap_csv.py → kaggle_upload/IEMOCAP/ (train/dev/test.csv)"),
        ],
        "total_time": "~19h",
        "summary_tasks": """-Key tasks done:
1. Download và xác minh cả 3 datasets: MELD (13,708 utts), IEMOCAP (7,380 utts), DailyDialog (102,979 utts)
2. Implement data loaders thống nhất cho 3 datasets với interface Dialogue/Utterance chung
3. Implement Dirichlet non-IID federated partitioning cho N clients
4. Extract frozen RoBERTa-base embeddings (768-dim) cho toàn bộ datasets
5. Export IEMOCAP sang CSV format để hỗ trợ Kaggle fine-tuning workflow
6. Xác minh class distribution và data quality cho cả 3 datasets""",
        "summary_learned": """-Key things learned:
1. IEMOCAP dataset structure: Session-based organization cần mapping phức tạp (session → dialogue → utterance)
2. Non-IID data partitioning: Dirichlet(alpha=0.5) tạo phân phối heterogeneous giống thực tế
3. Feature extraction pipeline: Pre-extracted embeddings giảm training time đáng kể (x10 faster)
4. DailyDialog imbalance: >80% utterances là "no_emotion" — cần xử lý đặc biệt khi evaluate
5. Data preprocessing: CSV export standardizes data format across different dataset structures""",
        "summary_literature": """-Literature read:
1. Busso et al. (2008) "IEMOCAP" — Dataset collection methodology, emotion annotation protocol
2. Poria et al. (2019) "MELD" — Multimodal EmotionLines Dataset from Friends TV series
3. Li et al. (2017) "DailyDialog" — Human-written daily conversation dataset
4. Hsu et al. (2022) "Non-IID Federated Learning" — Dirichlet partitioning for realistic FL simulation
5. Liu et al. (2019) "RoBERTa: A Robustly Optimized BERT Pretraining Approach" — Feature extraction backbone""",
        "summary_issues": """-Issues/Challenges:
1. IEMOCAP raw data structure phức tạp — nhiều file format (.txt, .wav, .avi) cần custom parser
2. IEMOCAP consensus annotation parsing — cần xử lý multiple annotator labels
3. DailyDialog no_emotion class chiếm >80% — evaluation protocol phải exclude hoặc handle riêng
4. Feature file size lớn (DailyDialog 311MB) — cần git-lfs hoặc separate storage
5. Cross-dataset standardization: mỗi dataset có emotion taxonomy khác nhau""",
        "next_plan": """Plan for the next week:
1. Extract WavLM audio features cho MELD (multimodal experiments)
2. Hỗ trợ viết paper: Data Description section
3. Verify ablation experiment data consistency""",
        "next_tasks": [
            ("Extract WavLM audio features cho MELD", "Week 2 (May 20)"),
            ("Viết paper: Dataset Description + Experimental Setup", "Week 2 (May 22)"),
            ("Data quality check cho ablation experiments", "Week 2 (May 19)"),
        ],
    },
    {
        "name": "Hồ Gia Phú",
        "id": "22110060",
        "filename": r"d:\OJT\Worklog - W1 - 22110060 - HoGiaPhu.docx",
        "tasks": [
            ("Implement DialogueRNN backbone: 3-GRU architecture (Global, Speaker, Emotion state tracking)",
             "Completed", "5h",
             "models/erc/dialogue_rnn.py — input 768-dim → context-aware hidden states"),
            ("Implement FL infrastructure: FederatedClient (local training) + FederatedServer (orchestration)",
             "Completed", "4h",
             "federated/client.py, server.py — support EDL loss, FedProx proximal term"),
            ("Implement FL aggregation: FedAvg (weighted average) + FedProx (proximal regularization)",
             "Completed", "3h",
             "federated/aggregation/fedavg.py, fedprox.py — standard FL baselines"),
            ("Implement FixMatch SSL: pseudo-labeling, curriculum threshold, feature augmentation",
             "Completed", "4h",
             "semi_supervised/fixmatch.py, augmentation.py — Gaussian noise + feature dropout"),
            ("Implement Dempster-Shafer fusion module cho multimodal evidence combination",
             "Completed", "3h",
             "models/evidential/ds_fusion.py — text+audio uncertainty-aware fusion"),
            ("Train centralized DialogueRNN+EDL baseline trên MELD và DailyDialog",
             "Completed", "2h",
             "MELD EDL WF1=0.6262, DailyDialog EDL WF1=0.8773 (seed 42)"),
        ],
        "total_time": "~21h",
        "summary_tasks": """-Key tasks done:
1. Implement DialogueRNN backbone model với 3-GRU architecture cho emotion tracking trong conversations
2. Implement complete Federated Learning infrastructure: Client (local training) + Server (global coordination)
3. Implement FedAvg + FedProx aggregation strategies làm baseline comparison
4. Implement FixMatch Semi-Supervised Learning module với feature-level augmentation
5. Implement Dempster-Shafer fusion cho multimodal evidence combination
6. Validate centralized training pipeline: EDL achieves WF1=0.6262 on MELD test set""",
        "summary_learned": """-Key things learned:
1. DialogueRNN: 3-GRU design (global/speaker/emotion) captures conversational context better than flat classifiers
2. Federated Learning simulation: Client-server architecture with parameter broadcasting and local SGD
3. FedProx: Proximal term (mu=0.01) prevents client drift in non-IID settings
4. Semi-supervised challenges in ERC: FixMatch's confidence thresholding struggles with dialogue-level predictions
5. Dempster-Shafer theory: Combining evidence from multiple modalities using belief functions""",
        "summary_literature": """-Literature read:
1. Majumder et al. (2019) "DialogueRNN: An Attentive RNN for Emotion Detection in Conversations"
2. Li et al. (2020) "Federated Optimization in Heterogeneous Networks" — FedProx algorithm
3. Sohn et al. (2020) "FixMatch: Simplifying Semi-Supervised Learning" — Pseudo-labeling with augmentation
4. Han et al. (2021) "Trusted Multi-View Classification with Dempster-Shafer Combination"
5. Shafer (1976) "A Mathematical Theory of Evidence" — Foundation for DS Fusion""",
        "summary_issues": """-Issues/Challenges:
1. DialogueRNN gradient vanishing trên dialogues dài (>30 utterances) — cần gradient clipping
2. FixMatch produces 0 pseudo-labels initially — threshold too high → added curriculum warmup
3. SSL 20% labels (WF1=0.4141) underperforms supervised (WF1=0.4736) — Poison Amplification effect
4. DS Fusion chỉ test trên frozen features — cần update với fine-tuned features
5. FedProx performance (WF1=0.5412) gần bằng FedAvg (WF1=0.5419) — mu sensitivity analysis needed""",
        "next_plan": """Plan for the next week:
1. Chạy ablation study: isolate EDL vs Softmax, EAFA vs FedAvg contribution
2. Hỗ trợ viết paper: Methodology section (DialogueRNN + EDL + EAFA)
3. Update DS Fusion với fine-tuned features""",
        "next_tasks": [
            ("Ablation: Softmax CE baseline (thay EDL bằng standard cross-entropy)", "Week 2 (May 20)"),
            ("Viết paper: Methodology — DialogueRNN + EDL + EAFA architecture", "Week 2 (May 22)"),
            ("Update DS Fusion với fine-tuned features cho multimodal", "Week 2 (May 22)"),
        ],
    },
]


def generate_worklog(member):
    doc = docx.Document(TEMPLATE)

    # Table 0: Project Info
    info_table = doc.tables[0]
    info_table.rows[0].cells[1].text = PROJECT_NAME
    info_table.rows[1].cells[1].text = member["name"]
    info_table.rows[2].cells[1].text = member["id"]

    # Table 1: Weekly Tasks
    task_table = doc.tables[1]

    # Delete data rows (keep header rows 0,1,2)
    while len(task_table.rows) > 3:
        tr = task_table.rows[3]._tr
        task_table._tbl.remove(tr)

    # Add task rows
    for task, status, time_spent, note in member["tasks"]:
        row = task_table.add_row()
        row.cells[0].text = task
        row.cells[1].text = status
        row.cells[2].text = time_spent
        row.cells[3].text = note

    # Total row
    total_row = task_table.add_row()
    total_row.cells[0].text = "TOTAL WEEKLY TIME SPENT"
    total_row.cells[1].text = ""
    total_row.cells[2].text = member["total_time"]
    total_row.cells[3].text = ""

    # Dates
    task_table.rows[0].cells[1].text = "1"
    task_table.rows[0].cells[2].text = "1"
    task_table.rows[0].cells[3].text = "1"
    task_table.rows[1].cells[1].text = "09/05/2026 - 15/05/2026"
    task_table.rows[1].cells[2].text = ""
    task_table.rows[1].cells[3].text = ""

    # Paragraphs
    doc.paragraphs[6].text = member["summary_tasks"]
    doc.paragraphs[8].text = member["summary_learned"]
    doc.paragraphs[10].text = member["summary_literature"]
    doc.paragraphs[12].text = member["summary_issues"]
    doc.paragraphs[15].text = member["next_plan"]

    # Table 2: Next Week Plan
    plan_table = doc.tables[2]
    while len(plan_table.rows) > 1:
        tr = plan_table.rows[1]._tr
        plan_table._tbl.remove(tr)

    for task, expected in member["next_tasks"]:
        row = plan_table.add_row()
        row.cells[0].text = task
        row.cells[1].text = expected

    doc.save(member["filename"])
    print(f"  Saved: {member['filename']}")


if __name__ == "__main__":
    print("Generating Work Logs for Week 1...")
    for m in members:
        generate_worklog(m)
    print("\nDone! 3 worklogs generated.")
